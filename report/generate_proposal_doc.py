import os
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, PageBreak
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
REPORT_DIR = os.path.join(BASE, "report")
DOCX_PATH = os.path.join(REPORT_DIR, "project_proposal_blueprint.docx")
PDF_PATH = os.path.join(REPORT_DIR, "project_proposal_blueprint.pdf")

# -----------------------------------------------------------------------------
# 1. WORD (.DOCX) PROPOSAL GENERATION
# -----------------------------------------------------------------------------
doc = Document()
for s in doc.sections:
    s.top_margin = Inches(0.8)
    s.bottom_margin = Inches(0.8)
    s.left_margin = Inches(0.8)
    s.right_margin = Inches(0.8)

PRIMARY = RGBColor(15, 23, 42)
SECONDARY = RGBColor(37, 99, 235)
HEX_PRIMARY = "0F172A"
HEX_BG = "F8FAFC"

def set_cell_bg(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'))

# Title
p_t = doc.add_paragraph()
r_t = p_t.add_run("Project Proposal & Implementation Blueprint")
r_t.font.name = "Calibri"
r_t.font.size = Pt(22)
r_t.font.bold = True
r_t.font.color.rgb = PRIMARY

p_sub = doc.add_paragraph()
r_sub = p_sub.add_run("Synthetic Identity Fraud Detection in Digital Lending Onboarding (KYC + Behavioral Signals)")
r_sub.font.name = "Calibri"
r_sub.font.size = Pt(13)
r_sub.font.bold = True
r_sub.font.color.rgb = SECONDARY

p_meta = doc.add_paragraph()
p_meta.add_run("Domain: BFSI / Applied Machine Learning  |  GitHub: https://github.com/sgsatpute/bfsi").italic = True

doc.add_heading("1. Problem Statement", level=1)
doc.add_paragraph(
    "Digital lending platforms enable rapid loan approvals through automated customer onboarding. However, this automation introduces severe exposure to "
    "Synthetic Identity Fraud (SIF)—where fraudsters construct composite identities combining real PII fragments (e.g. valid PAN tax IDs) with fabricated contact details "
    "(burner phones, synthetic emails, false addresses):\n\n"
    "Synthetic Identity = P_real ∪ P_fabricated\n\n"
    "Why Rule-Based KYC Fails:\n"
    "1. Field-Level Database Match: Traditional KYC verifies isolated fields (OCR + database query). Because individual PII fields are valid, static queries return clean matches.\n"
    "2. Absence of Immediate Victims: Unlike stolen identity fraud, there is no real victim receiving charge alerts, allowing fraudsters to build credit scores before maxing out loan limits ('credit bust-out') and abandoning the account.\n\n"
    "Proposed Machine Learning Solution:\n"
    "We implement a multimodal machine learning framework unifying KYC entity matching, identity freshness metrics, behavioral biometrics (keystroke cadence, hesitation, paste ratio), "
    "and device/network velocity into a real-time risk engine. The engine computes a fraud probability score and routes applications into Automated Low (<30%), Medium (30-60%), and High (>60%) risk decision bands."
)

doc.add_heading("2. Team Members & Task Allocation Matrix", level=1)
team_table = doc.add_table(rows=5, cols=3)
team_table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Role / Area", "Key Responsibilities", "Deliverables Created"]
for j, h in enumerate(headers):
    cell = team_table.cell(0, j)
    cell.text = h
    set_cell_bg(cell, HEX_PRIMARY)
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)

t_rows = [
    ["1. Lead AI/ML Engineer", "Model selection, 5-fold Stratified CV, hyperparameter tuning, metric saving.", "train_model.py, fraud_model.joblib, metrics.json"],
    ["2. Data & Feature Engineer", "20-signal synthetic dataset design (10,000 apps), noise modeling, distributions.", "generate_data.py, synthetic_kyc_behavioral.csv"],
    ["3. Full-Stack / UI Engineer", "Streamlit research portal, live risk inference engine, batch telemetry inspector.", "dashboard.py (Streamlit Web Portal)"],
    ["4. Research & Doc Lead", "Literature survey, formal IEEE paper drafting, 12-slide presentation deck, charts.", "research_paper.pdf, presentation.pptx, roc_curves.png"]
]

for i, row in enumerate(t_rows):
    for j, val in enumerate(row):
        cell = team_table.cell(i+1, j)
        cell.text = val
        if i % 2 == 1:
            set_cell_bg(cell, HEX_BG)

doc.add_paragraph()

doc.add_heading("3. Implementation Roadmap & Blueprint", level=1)
roadmap_table = doc.add_table(rows=6, cols=3)
roadmap_table.alignment = WD_TABLE_ALIGNMENT.CENTER

r_headers = ["Phase & Timeline", "Key Objectives & Activities", "Phase Deliverable Output"]
for j, h in enumerate(r_headers):
    cell = roadmap_table.cell(0, j)
    cell.text = h
    set_cell_bg(cell, HEX_PRIMARY)
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)

r_data = [
    ["Phase 1 (Week 1)", "Problem definition, literature survey, threat taxonomy formalization.", "Problem Statement & Threat Model"],
    ["Phase 2 (Week 1–2)", "20-signal multimodal feature engineering, dataset generation (10,000 rows).", "synthetic_kyc_behavioral.csv"],
    ["Phase 3 (Week 2–3)", "5-model benchmark suite (LR, DT, RF, HistGBDT, MLP) & 5-fold CV evaluation.", "train_model.py, research_metrics.json"],
    ["Phase 4 (Week 3–4)", "Interactive Streamlit web portal development & real-time risk engine.", "dashboard.py (Active Port 8501)"],
    ["Phase 5 (Week 4)", "Formal IEEE research paper compilation, slide deck, and final defense.", "research_paper.pdf, presentation.pptx, ZIP archive"]
]

for i, row in enumerate(r_data):
    for j, val in enumerate(row):
        cell = roadmap_table.cell(i+1, j)
        cell.text = val
        if i % 2 == 1:
            set_cell_bg(cell, HEX_BG)

doc.add_paragraph()

doc.add_heading("4. Expected Deliverables & Benchmark Summary", level=1)
doc.add_paragraph(
    "Expected Deliverables:\n"
    "1. Complete Source Code & Repository: https://github.com/sgsatpute/bfsi\n"
    "2. IEEE Research Paper: report/research_paper.pdf & research_paper.docx\n"
    "3. Presentation Slide Deck: report/presentation.pptx & presentation.pdf\n"
    "4. Web Portal Demo: Interactive Streamlit application (app/dashboard.py)\n"
    "5. Submission ZIP Package: Synthetic_Identity_Fraud_Detection_Research_Level_Submission.zip\n\n"
    "Benchmarked Model Results (5-Fold Stratified Cross-Validation):\n"
    "• ROC-AUC: 0.9139\n"
    "• Fraud Precision: 0.8600\n"
    "• Fraud Recall: 0.9110\n"
    "• Fraud F1-Score: 0.8849"
)

doc.save(DOCX_PATH)
print("Generated Proposal DOCX at:", DOCX_PATH)

# -----------------------------------------------------------------------------
# 2. REPORTLAB PDF PROPOSAL GENERATION
# -----------------------------------------------------------------------------
pdf_doc = SimpleDocTemplate(
    PDF_PATH, pagesize=letter,
    rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40
)
styles = getSampleStyleSheet()

title_s = ParagraphStyle('PTitle', fontName='Helvetica-Bold', fontSize=18, leading=22, textColor=colors.HexColor('#0F172A'), spaceAfter=4)
sub_s = ParagraphStyle('PSub', fontName='Helvetica-Bold', fontSize=11, leading=14, textColor=colors.HexColor('#2563EB'), spaceAfter=10)
h1_s = ParagraphStyle('PH1', fontName='Helvetica-Bold', fontSize=12, leading=15, textColor=colors.HexColor('#0F172A'), spaceBefore=10, spaceAfter=4, keepWithNext=True)
body_s = ParagraphStyle('PBody', fontName='Helvetica', fontSize=9, leading=13, textColor=colors.HexColor('#1E293B'), alignment=TA_JUSTIFY, spaceAfter=6)
th_s = ParagraphStyle('PTH', fontName='Helvetica-Bold', fontSize=8, leading=10, textColor=colors.white)
tb_s = ParagraphStyle('PTB', fontName='Helvetica', fontSize=8, leading=10, textColor=colors.HexColor('#1E293B'))

story = []
story.append(Paragraph("Project Proposal & Implementation Blueprint", title_s))
story.append(Paragraph("Synthetic Identity Fraud Detection in Digital Lending Onboarding", sub_s))
story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#2563EB'), spaceAfter=10))

story.append(Paragraph("1. Problem Statement Overview", h1_s))
story.append(Paragraph(
    "Digital lending platforms enable rapid loan approvals through automated customer onboarding. However, this introduces severe exposure to "
    "<b>Synthetic Identity Fraud (SIF)</b>—where fraudsters construct composite identities combining real PII fragments with fabricated contact details.<br/><br/>"
    "<b>Why Static KYC Fails:</b> Individual PII fields pass isolated OCR/database checks. Furthermore, absence of immediate victims allows fraudsters to cultivate credit scores before maxing out credit limits.<br/><br/>"
    "<b>ML Solution:</b> We implement a 20-signal multimodal ML framework unifying KYC matching, identity age, behavioral biometrics (keystroke cadence, paste ratio), and device velocity into a real-time risk engine.",
    body_s
))

story.append(Paragraph("2. Team Roles & Task Allocation", h1_s))
t_pdf_data = [
    [Paragraph("Role / Area", th_s), Paragraph("Key Responsibilities", th_s), Paragraph("Deliverables Created", th_s)]
] + [[Paragraph(r[0], tb_s), Paragraph(r[1], tb_s), Paragraph(r[2], tb_s)] for r in t_rows]

t_pdf = Table(t_pdf_data, colWidths=[120, 220, 190])
t_pdf.setStyle(TableStyle([
    ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#0F172A')),
    ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
    ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#F8FAFC')]),
    ('TOPPADDING', (0,0), (-1,-1), 4), ('BOTTOMPADDING', (0,0), (-1,-1), 4),
]))
story.append(t_pdf)

story.append(Paragraph("3. Implementation Roadmap", h1_s))
r_pdf_data = [
    [Paragraph("Phase & Timeline", th_s), Paragraph("Objectives & Activities", th_s), Paragraph("Phase Output", th_s)]
] + [[Paragraph(r[0], tb_s), Paragraph(r[1], tb_s), Paragraph(r[2], tb_s)] for r in r_data]

r_pdf = Table(r_pdf_data, colWidths=[110, 240, 180])
r_pdf.setStyle(TableStyle([
    ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#0F172A')),
    ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
    ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#F8FAFC')]),
    ('TOPPADDING', (0,0), (-1,-1), 4), ('BOTTOMPADDING', (0,0), (-1,-1), 4),
]))
story.append(r_pdf)

story.append(Paragraph("4. Benchmark Metrics & Repository", h1_s))
story.append(Paragraph(
    "<b>GitHub Repository:</b> https://github.com/sgsatpute/bfsi<br/>"
    "<b>5-Fold Cross-Validation Metrics:</b> ROC-AUC: <b>0.9139</b> | Fraud Precision: <b>0.8600</b> | Fraud Recall: <b>0.9110</b> | F1-Score: <b>0.8849</b><br/>"
    "<b>Deliverables:</b> Source Code, 10k Dataset, IEEE Research Paper (PDF/DOCX), 12-Slide PPTX Presentation, Streamlit Web Portal, Submission ZIP Package.",
    body_s
))

pdf_doc.build(story)
print("Generated Proposal PDF at:", PDF_PATH)
