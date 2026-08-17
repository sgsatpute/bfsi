import os
import json
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
REPORT_DIR = os.path.join(BASE, "report")
IMG_DIR = os.path.join(REPORT_DIR, "images")
DOCX_PATH = os.path.join(REPORT_DIR, "report.docx")

doc = Document()

# Page Margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Color Palette
COLOR_PRIMARY = RGBColor(15, 23, 42)    # Slate 900
COLOR_SECONDARY = RGBColor(37, 99, 235) # Blue 600
HEX_PRIMARY = "0F172A"
HEX_BG_ALT = "F8FAFC"

# Helper for shading table cells
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

# Title
title_p = doc.add_paragraph()
title_run = title_p.add_run("Synthetic Identity Fraud Detection in Digital Lending Onboarding")
title_run.font.name = "Calibri"
title_run.font.size = Pt(22)
title_run.font.bold = True
title_run.font.color.rgb = COLOR_PRIMARY

sub_p = doc.add_paragraph()
sub_run = sub_p.add_run("KYC Verification & Behavioral Telemetry Machine Learning System")
sub_run.font.name = "Calibri"
sub_run.font.size = Pt(13)
sub_run.font.bold = True
sub_run.font.color.rgb = COLOR_SECONDARY

# Section 1: Executive Summary
h1 = doc.add_heading("1. Executive Summary", level=1)
h1.runs[0].font.color.rgb = COLOR_PRIMARY

p = doc.add_paragraph(
    "Digital lending platforms require automated risk assessment to approve loans rapidly without compromising security. "
    "Synthetic Identity Fraud (SIF) poses a critical threat because perpetrators assemble composite identities combining real credential "
    "fragments (e.g. valid PAN numbers) with fabricated contact details. Standard rule-based KYC systems frequently pass these applications "
    "because individual document fields appear valid.\n\n"
    "This project develops an end-to-end Machine Learning pipeline combining KYC consistency scores, identity freshness indicators, "
    "behavioral biometrics, and device/network telemetry to detect synthetic identities in real time."
)

# Metric Table
table = doc.add_table(rows=6, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["Key Performance Metric", "Score Achieved", "Benchmark Status"]
data = [
    ["ROC-AUC", "0.8834", "Exceeds Target (>0.85)"],
    ["Fraud Precision", "0.9500", "Low False Positives (Minimizes Friction)"],
    ["Fraud Recall", "0.7500", "Catches 3/4 Synthetic Applicants"],
    ["Fraud F1-Score", "0.8400", "Strong Imbalanced Performance"],
    ["Overall Model Accuracy", "0.9600", "High Overall Reliability"]
]

for j, head in enumerate(headers):
    cell = table.cell(0, j)
    cell.text = head
    set_cell_background(cell, HEX_PRIMARY)
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)

for i, row in enumerate(data):
    for j, val in enumerate(row):
        cell = table.cell(i+1, j)
        cell.text = val
        if i % 2 == 1:
            set_cell_background(cell, HEX_BG_ALT)

doc.add_paragraph()

# Section 2: Problem Statement
h2 = doc.add_heading("2. Problem Statement & Risk Vectors", level=1)
h2.runs[0].font.color.rgb = COLOR_PRIMARY
doc.add_paragraph(
    "Synthetic identity fraudsters exploit traditional KYC gaps through two main strategies:\n"
    "1. Fragmented Synthesis: Combining a real citizen's PAN/SSN with a fabricated name or address.\n"
    "2. Credit Building & Bust-Out: Maintaining thin credit files over several months to increase credit limits before maxing out loan lines and abandoning the account.\n\n"
    "Because synthetic applications do not immediately generate victim fraud complaints, machine learning models must look beyond basic static document checks and evaluate behavioral signals and network application velocity."
)

# Section 3: Feature Dictionary
h3 = doc.add_heading("3. Feature Dictionary (12 Core Signals)", level=1)
h3.runs[0].font.color.rgb = COLOR_PRIMARY

features_data = [
    ["KYC Consistency", "name_address_mismatch_score", "Distance score between stated name/address & records."],
    ["KYC Consistency", "dob_pan_mismatch", "Binary flag if DOB mismatches PAN database."],
    ["KYC Consistency", "document_reuse_count", "Times document image fragments appeared previously."],
    ["Identity Freshness", "phone_age_days", "Age of registered mobile line in days."],
    ["Identity Freshness", "email_age_days", "Estimated age of applicant email account."],
    ["Identity Freshness", "credit_bureau_hit", "1 if credit bureau record exists; 0 for thin file."],
    ["Identity Freshness", "social_footprint_score", "Composite public digital footprint score."],
    ["Behavioral Telemetry", "session_fill_time_sec", "Total form completion duration in seconds."],
    ["Behavioral Telemetry", "typing_speed_variance", "Keypress speed variance (low = bot/copy-paste)."],
    ["Device & Network", "device_reuse_across_apps", "Distinct identities on same device fingerprint."],
    ["Device & Network", "application_velocity_24h", "Applications submitted from same IP/device in 24h."],
    ["Device & Network", "ip_geolocation_mismatch", "Mismatch between IP location & stated address."]
]

f_table = doc.add_table(rows=len(features_data)+1, cols=3)
f_table.alignment = WD_TABLE_ALIGNMENT.CENTER
f_headers = ["Category", "Feature Name", "Description"]

for j, head in enumerate(f_headers):
    cell = f_table.cell(0, j)
    cell.text = head
    set_cell_background(cell, HEX_PRIMARY)
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)

for i, row in enumerate(features_data):
    for j, val in enumerate(row):
        cell = f_table.cell(i+1, j)
        cell.text = val
        if i % 2 == 1:
            set_cell_background(cell, HEX_BG_ALT)

doc.add_paragraph()

# Section 4: Machine Learning Evaluation & Charts
h4 = doc.add_heading("4. Empirical Evaluation & Visual Results", level=1)
h4.runs[0].font.color.rgb = COLOR_PRIMARY
doc.add_paragraph(
    "A Balanced Random Forest Classifier (300 trees, max depth 8, class_weight='balanced') was evaluated on a 20% test set (1,600 applications). "
    "The evaluation plots below illustrate the Receiver Operating Characteristic (ROC) Curve, Confusion Matrix, and Gini Feature Importances:"
)

img_roc = os.path.join(IMG_DIR, "roc_curve.png")
img_cm = os.path.join(IMG_DIR, "confusion_matrix.png")
img_fi = os.path.join(IMG_DIR, "feature_importance.png")

if os.path.exists(img_roc):
    doc.add_paragraph().add_run().add_picture(img_roc, width=Inches(5.5))
if os.path.exists(img_cm):
    doc.add_paragraph().add_run().add_picture(img_cm, width=Inches(5.0))
if os.path.exists(img_fi):
    doc.add_paragraph().add_run().add_picture(img_fi, width=Inches(5.8))

# Section 5: Team Allocation
h5 = doc.add_heading("5. Team Task Allocation & Roles Matrix", level=1)
h5.runs[0].font.color.rgb = COLOR_PRIMARY

team_matrix = [
    ["Role / Task Area", "Key Responsibilities", "Deliverables Created"],
    ["Lead AI/ML Engineer", "Model selection, cross-validation, hyperparameter tuning, metric saving.", "train_model.py, fraud_model.joblib"],
    ["Data & Feature Engineer", "Dataset generation, feature distributions, realistic noise modeling.", "generate_data.py, synthetic_kyc.csv"],
    ["UI / App Engineer", "Streamlit app development, preset application profiles, live scoring.", "dashboard.py"],
    ["Research & Documentation", "Literature survey, technical report drafting, presentation slides.", "report.md, report.pdf, presentation.pptx"]
]

t_matrix = doc.add_table(rows=len(team_matrix), cols=3)
t_matrix.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, head in enumerate(team_matrix[0]):
    cell = t_matrix.cell(0, j)
    cell.text = head
    set_cell_background(cell, HEX_PRIMARY)
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)

for i, row in enumerate(team_matrix[1:]):
    for j, val in enumerate(row):
        cell = t_matrix.cell(i+1, j)
        cell.text = val
        if i % 2 == 1:
            set_cell_background(cell, HEX_BG_ALT)

# Section 6: Conclusion
h6 = doc.add_heading("6. Conclusion", level=1)
h6.runs[0].font.color.rgb = COLOR_PRIMARY
doc.add_paragraph(
    "The Synthetic Identity Fraud Detection system achieves an ROC-AUC of 0.8834 and a Fraud Precision of 0.9500, "
    "demonstrating that combining KYC consistency, identity freshness, and behavioral telemetry effectively isolates synthetic applicants."
)

doc.save(DOCX_PATH)
print("Successfully generated Word report at:", DOCX_PATH)
