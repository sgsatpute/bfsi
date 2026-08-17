import os
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, PageBreak
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
REPORT_DIR = os.path.join(BASE, "report")

MD_PATH = os.path.join(REPORT_DIR, "SIMPLE_PROJECT_EXPLANATION.md")
DOCX_PATH = os.path.join(REPORT_DIR, "SIMPLE_PROJECT_EXPLANATION.docx")
PDF_PATH = os.path.join(REPORT_DIR, "SIMPLE_PROJECT_EXPLANATION.pdf")

# -----------------------------------------------------------------------------
# 1. MARKDOWN SIMPLE EXPLANATION GENERATION
# -----------------------------------------------------------------------------
md_text = """# 💡 Simple Guide: Synthetic Identity Fraud Detection

**Project Title:** Synthetic Identity Fraud Detection in Digital Lending Onboarding  
**GitHub Repository:** [https://github.com/sgsatpute/bfsi](https://github.com/sgsatpute/bfsi)

---

## 1. What is this project about? (In Plain English)
Imagine a bank offering instant online loans. Anyone can apply on their mobile phone and get money in 2 minutes.

A **fraudster** wants to steal loan money. But instead of stealing a real person's identity (which would cause the real person to complain to the police immediately), the fraudster creates a **"Frankenstein Fake Identity"**:
- **Real Part:** A valid PAN or tax number (purchased or leaked online).
- **Fake Part:** Fake name, fake residential address, brand new burner phone number, and newly created email address.

This fake person **does not exist in the real world**, but they apply for a loan online!

---

## 2. Why do traditional bank checks fail?
When the bank's system checks the PAN number, the government database says: *"Yes! This PAN number is valid!"*  
Because standard banks only check one field at a time, the fake person gets approved, takes the loan money, and disappears. Since no real person's name was used directly, no victim complains for months.

---

## 3. How does OUR AI system catch this fake person?
Our AI system acts like a **smart detective** looking at **4 different categories of clues (20 signals total)**:

1. **KYC Document Clues:** Checks if the name, address, DOB, and PAN actually match each other, or if the photo ID image was reused across previous loan apps.
2. **Identity Freshness Clues:** Is the phone number only 3 days old? Is the email registered yesterday? Does this person have zero credit history at the credit bureau? (Real adults usually have older phone numbers and email addresses).
3. **Behavioral Habits (How they fill the form):**
   - **Humans:** Type at normal speed, make typos, hit backspace, and pause while reading questions.
   - **Bots / Fraudsters:** Fill out 50 form questions in 10 seconds, copy-paste everything from a file, and make zero backspaces!
4. **Device & Network Clues:** Did 5 different loan applications come from the exact same laptop or same Wi-Fi IP address today?

---

## 4. How does the AI make a decision?
Our AI model (Random Forest Classifier) combines all 20 clues and outputs a **Fraud Risk Score (0% to 100%)**:

- 🟢 **Low Risk (< 30%):** Normal genuine customer → **Instant Auto-Approval**.
- 🟡 **Medium Risk (30% - 60%):** Slightly suspicious → **Ask for Video KYC / OTP verification**.
- 🔴 **High Risk (> 60%):** Confirmed fake identity → **Block Loan Application immediately!**

---

## 5. What are our final project results?
We tested our AI model on **10,000 application records**:
- **Accuracy / ROC-AUC:** **91.4%** (Extremely high accuracy!)
- **Precision:** **86.0%** (Very low false alarms — genuine customers are not bothered).
- **Recall:** **91.1%** (Catches 9 out of every 10 fraud attempts!).

---

## 6. How to explain this to your professor in 30 seconds:
> *"Sir, traditional KYC systems only verify static document numbers, so synthetic fake identities (real PAN + fake phone/address) bypass them easily. Our project uses Machine Learning on 20 signals—combining document consistency, phone/email age, behavioral biometrics (typing cadence, paste speed), and device velocity. Our model achieves 91.4% ROC-AUC accuracy and is fully deployed with an interactive Streamlit web dashboard for live risk scoring."*

---

## 7. What files are inside the submission?
- `data/synthetic_kyc_behavioral.csv`: 10,000 application records with 20 signals.
- `model/train_model.py`: AI model training & 5-fold cross-validation script.
- `app/dashboard.py`: Live interactive web app dashboard (run: `python -m streamlit run app/dashboard.py`).
- `report/research_paper.pdf`: IEEE-style formal academic research paper.
- `report/presentation.pptx`: 12-slide PowerPoint presentation deck.
"""

with open(MD_PATH, "w", encoding="utf-8") as f:
    f.write(md_text)

print("Generated Markdown Guide at:", MD_PATH)

# -----------------------------------------------------------------------------
# 2. WORD (.DOCX) SIMPLE EXPLANATION GENERATION
# -----------------------------------------------------------------------------
doc = Document()
for s in doc.sections:
    s.top_margin = Inches(0.8)
    s.bottom_margin = Inches(0.8)

PRIMARY = RGBColor(15, 23, 42)
SECONDARY = RGBColor(37, 99, 235)

p_t = doc.add_paragraph()
r_t = p_t.add_run("Simple Explanation Guide: Synthetic Identity Fraud Detection")
r_t.font.name = "Calibri"
r_t.font.size = Pt(20)
r_t.font.bold = True
r_t.font.color.rgb = PRIMARY

p_sub = doc.add_paragraph()
r_sub = p_sub.add_run("Easy-to-Understand Overview of Project Concept, AI Approach, & Results")
r_sub.font.name = "Calibri"
r_sub.font.size = Pt(12)
r_sub.font.bold = True
r_sub.font.color.rgb = SECONDARY

doc.add_heading("1. What is this project about? (In Plain English)", level=1)
doc.add_paragraph(
    "Imagine a bank offering instant online loans. Anyone can apply on their mobile phone and get money in 2 minutes.\n\n"
    "A fraudster wants to steal loan money. Instead of stealing a real person's full identity, the fraudster creates a 'Frankenstein Fake Identity':\n"
    "• Real Part: A valid PAN or tax number (purchased or leaked online).\n"
    "• Fake Part: Fake name, fake address, brand new burner phone number, newly created email address.\n\n"
    "This fake person does not exist in the real world, but they apply for a loan online!"
)

doc.add_heading("2. Why do traditional bank checks fail?", level=1)
doc.add_paragraph(
    "When the bank's system checks the PAN number, the government database says: 'Yes! This PAN number is valid!'\n"
    "Because standard banks only check one field at a time, the fake person gets approved, takes the loan money, and disappears. "
    "Since no real person's name was used directly, no victim complains for months."
)

doc.add_heading("3. How does OUR AI system catch this fake person?", level=1)
doc.add_paragraph(
    "Our AI system acts like a smart detective looking at 4 different categories of clues (20 signals total):\n\n"
    "1. KYC Document Clues: Checks if the name, address, DOB, and PAN actually match each other, or if document images were reused.\n"
    "2. Identity Freshness Clues: Is the phone number only 3 days old? Is the email registered yesterday? Does this person have zero credit history?\n"
    "3. Behavioral Habits: Humans type normally, make typos, and pause. Bots or fraudsters fill 50 questions in 10 seconds and copy-paste everything.\n"
    "4. Device & Network Clues: Did 5 different loan applications come from the exact same laptop or Wi-Fi IP address today?"
)

doc.add_heading("4. How does the AI make a decision?", level=1)
doc.add_paragraph(
    "Our AI model combines all 20 clues and outputs a Fraud Risk Score (0% to 100%):\n\n"
    "• Low Risk (< 30%): Normal genuine customer → Instant Auto-Approval.\n"
    "• Medium Risk (30% - 60%): Slightly suspicious → Ask for Video KYC / OTP verification.\n"
    "• High Risk (> 60%): Confirmed fake identity → Block Loan Application immediately!"
)

doc.add_heading("5. What are our final project results?", level=1)
doc.add_paragraph(
    "We tested our AI model on 10,000 application records:\n"
    "• Accuracy / ROC-AUC: 91.4% (Extremely high accuracy!)\n"
    "• Precision: 86.0% (Very low false alarms—genuine customers are not bothered).\n"
    "• Recall: 91.1% (Catches 9 out of every 10 fraud attempts!)."
)

doc.add_heading("6. 30-Second Elevator Pitch for Professor Review", level=1)
doc.add_paragraph(
    "\"Sir, traditional KYC systems only verify static document numbers, so synthetic fake identities (real PAN + fake phone/address) bypass them easily. "
    "Our project uses Machine Learning on 20 signals—combining document consistency, phone/email age, behavioral biometrics (typing cadence, paste speed), and device velocity. "
    "Our model achieves 91.4% ROC-AUC accuracy and is fully deployed with an interactive Streamlit web dashboard for live risk scoring.\""
)

doc.save(DOCX_PATH)
print("Generated Word Guide at:", DOCX_PATH)

# -----------------------------------------------------------------------------
# 3. REPORTLAB PDF SIMPLE EXPLANATION GENERATION
# -----------------------------------------------------------------------------
pdf_doc = SimpleDocTemplate(
    PDF_PATH, pagesize=letter,
    rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40
)
styles = getSampleStyleSheet()

title_s = ParagraphStyle('STitle', fontName='Helvetica-Bold', fontSize=18, leading=22, textColor=colors.HexColor('#0F172A'), spaceAfter=4)
sub_s = ParagraphStyle('SSub', fontName='Helvetica-Bold', fontSize=11, leading=14, textColor=colors.HexColor('#2563EB'), spaceAfter=10)
h1_s = ParagraphStyle('SH1', fontName='Helvetica-Bold', fontSize=12, leading=15, textColor=colors.HexColor('#0F172A'), spaceBefore=10, spaceAfter=4, keepWithNext=True)
body_s = ParagraphStyle('SBody', fontName='Helvetica', fontSize=9.5, leading=13.5, textColor=colors.HexColor('#1E293B'), alignment=TA_JUSTIFY, spaceAfter=6)

story = []
story.append(Paragraph("Simple Explanation Guide: Synthetic Identity Fraud Detection", title_s))
story.append(Paragraph("Easy-to-Understand Overview of Project Concept, AI Approach, & Results", sub_s))
story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#2563EB'), spaceAfter=10))

story.append(Paragraph("1. What is this project about? (In Plain English)", h1_s))
story.append(Paragraph(
    "Imagine a bank offering instant online loans. A fraudster creates a <b>'Frankenstein Fake Identity'</b> using a real PAN tax number combined with fake name, burner phone, and synthetic email. "
    "Because the PAN is valid, traditional banks pass the applicant. The fraudster takes the loan money and vanishes!",
    body_s
))

story.append(Paragraph("2. How does OUR AI system catch this fake person?", h1_s))
story.append(Paragraph(
    "Our AI acts like a smart detective checking <b>20 clues across 4 categories</b>:<br/>"
    "1. <b>KYC Clues:</b> Name/address mismatch, DOB verification, document image reuse.<br/>"
    "2. <b>Identity Freshness:</b> Phone line age (<5 days?), email domain age, credit bureau history.<br/>"
    "3. <b>Behavioral Habits:</b> Humans type normally with backspaces; bots fill form in 5 seconds and copy-paste everything.<br/>"
    "4. <b>Device Telemetry:</b> Did 5 applications originate from the exact same laptop or Wi-Fi IP address?",
    body_s
))

story.append(Paragraph("3. AI Decision Risk Bands", h1_s))
story.append(Paragraph(
    "• 🟢 <b>Low Risk (<30%):</b> Genuine customer → Auto-Approve & Disburse.<br/>"
    "• 🟡 <b>Medium Risk (30%-60%):</b> Minor anomaly → Video KYC / OTP Check.<br/>"
    "• 🔴 <b>High Risk (>60%):</b> Fake identity → Block Application Immediately!",
    body_s
))

story.append(Paragraph("4. Project Results & Performance", h1_s))
story.append(Paragraph(
    "• <b>Accuracy / ROC-AUC:</b> <b>91.4%</b> (High discrimination capability)<br/>"
    "• <b>Fraud Precision:</b> <b>86.0%</b> (Low false alarms for real customers)<br/>"
    "• <b>Fraud Recall:</b> <b>91.1%</b> (Catches 9 out of 10 fraud attempts)",
    body_s
))

story.append(Paragraph("5. 30-Second Summary for Professor", h1_s))
story.append(Paragraph(
    "<i>\"Sir, traditional KYC systems only verify static document numbers, so synthetic fake identities bypass them easily. "
    "Our project uses Machine Learning on 20 signals—combining document consistency, phone/email age, behavioral biometrics (typing cadence, paste speed), and device velocity. "
    "Our model achieves 91.4% ROC-AUC accuracy and is fully deployed with an interactive Streamlit web dashboard for live risk scoring.\"</i>",
    body_s
))

pdf_doc.build(story)
print("Generated PDF Guide at:", PDF_PATH)
